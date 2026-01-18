"""
Document Parser Service for MathPulse AI
Intelligently parses various file formats (CSV, Excel, PDF) and extracts student data
Uses Hugging Face models for understanding unstructured/non-standard formats
"""
import io
import re
import json
import traceback
from typing import Optional, List, Dict, Any, Tuple
from dataclasses import dataclass
from enum import Enum
import chardet

# Data processing
import pandas as pd
import numpy as np

# PDF processing
try:
    import pdfplumber
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    print("Warning: pdfplumber not installed. PDF support disabled.")

try:
    from PyPDF2 import PdfReader
    PYPDF2_SUPPORT = True
except ImportError:
    PYPDF2_SUPPORT = False

# Word document processing
try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    print("Warning: python-docx not installed. DOCX support disabled.")


class FileType(str, Enum):
    CSV = "csv"
    EXCEL = "excel"
    PDF = "pdf"
    DOCX = "docx"
    UNKNOWN = "unknown"


@dataclass
class ParsedStudent:
    """Represents a parsed student record"""
    name: str
    engagement_score: Optional[float] = None
    avg_quiz_score: Optional[float] = None
    attendance: Optional[float] = None
    weakest_topic: Optional[str] = None
    grades: Optional[Dict[str, float]] = None
    raw_data: Optional[Dict[str, Any]] = None
    
    def to_dict(self) -> dict:
        return {
            "name": self.name,
            "engagementScore": self.engagement_score or 50,
            "avgQuizScore": self.avg_quiz_score or 50,
            "attendance": self.attendance,
            "weakestTopic": self.weakest_topic or "General Mathematics",
            "grades": self.grades or {},
            "rawData": self.raw_data or {}
        }


@dataclass
class ParseResult:
    """Result of document parsing"""
    success: bool
    students: List[ParsedStudent]
    file_type: FileType
    columns_detected: List[str]
    mapping_confidence: float
    warnings: List[str]
    raw_preview: Optional[str] = None
    course_info: Optional[Dict[str, Any]] = None
    
    def to_dict(self) -> dict:
        return {
            "success": self.success,
            "students": [s.to_dict() for s in self.students],
            "fileType": self.file_type.value,
            "columnsDetected": self.columns_detected,
            "mappingConfidence": self.mapping_confidence,
            "warnings": self.warnings,
            "rawPreview": self.raw_preview,
            "courseInfo": self.course_info,
            "studentCount": len(self.students)
        }


class ColumnMapper:
    """
    Intelligent column mapping using pattern matching and heuristics
    Maps arbitrary column names to our standard fields
    """
    
    # Common patterns for each field type
    NAME_PATTERNS = [
        r'student.*name', r'full.*name', r'name', r'learner', r'pupil',
        r'first.*name', r'last.*name', r'surname', r'given.*name',
        r'student', r'nombre', r'nom', r'姓名'
    ]
    
    QUIZ_PATTERNS = [
        r'quiz', r'test', r'exam', r'assessment', r'score', r'grade',
        r'average', r'avg', r'mark', r'result', r'performance',
        r'midterm', r'final', r'quiz.*score', r'test.*score'
    ]
    
    ENGAGEMENT_PATTERNS = [
        r'engagement', r'participation', r'activity', r'involvement',
        r'active', r'interaction', r'attend', r'contribution'
    ]
    
    ATTENDANCE_PATTERNS = [
        r'attendance', r'present', r'absent', r'attended', r'attend',
        r'days.*present', r'attendance.*rate', r'absences'
    ]
    
    TOPIC_PATTERNS = [
        r'weak', r'topic', r'subject', r'area', r'struggle', r'difficult',
        r'needs.*improvement', r'focus.*area', r'challenge'
    ]
    
    ID_PATTERNS = [
        r'id', r'student.*id', r'roll.*no', r'number', r'registration',
        r'enrollment', r'admission.*no'
    ]
    
    @classmethod
    def find_best_match(cls, columns: List[str], patterns: List[str]) -> Optional[str]:
        """Find the best matching column for given patterns"""
        columns_lower = [c.lower().strip() for c in columns]
        
        for pattern in patterns:
            for i, col in enumerate(columns_lower):
                if re.search(pattern, col, re.IGNORECASE):
                    return columns[i]
        return None
    
    @classmethod
    def map_columns(cls, columns: List[str]) -> Dict[str, Optional[str]]:
        """
        Map detected columns to standard fields
        Returns a mapping dictionary
        """
        mapping = {
            'name': cls.find_best_match(columns, cls.NAME_PATTERNS),
            'quiz_score': cls.find_best_match(columns, cls.QUIZ_PATTERNS),
            'engagement': cls.find_best_match(columns, cls.ENGAGEMENT_PATTERNS),
            'attendance': cls.find_best_match(columns, cls.ATTENDANCE_PATTERNS),
            'weak_topic': cls.find_best_match(columns, cls.TOPIC_PATTERNS),
            'student_id': cls.find_best_match(columns, cls.ID_PATTERNS),
        }
        
        # If no name column found, use first text column
        if not mapping['name']:
            for col in columns:
                if not any(re.search(p, col.lower()) for p in cls.QUIZ_PATTERNS + cls.ID_PATTERNS):
                    mapping['name'] = col
                    break
        
        return mapping
    
    @classmethod
    def calculate_confidence(cls, mapping: Dict[str, Optional[str]]) -> float:
        """Calculate confidence score for the mapping"""
        found = sum(1 for v in mapping.values() if v is not None)
        total = len(mapping)
        base_confidence = found / total
        
        # Boost if name is found
        if mapping['name']:
            base_confidence += 0.2
        
        # Boost if quiz score found
        if mapping['quiz_score']:
            base_confidence += 0.15
            
        return min(base_confidence, 1.0)


class DocumentParser:
    """Main document parsing service"""
    
    def __init__(self, hf_service=None):
        self.hf_service = hf_service
        self.column_mapper = ColumnMapper()
    
    def detect_file_type(self, filename: str, content: bytes) -> FileType:
        """Detect file type from filename and content"""
        filename_lower = filename.lower()
        
        if filename_lower.endswith('.csv'):
            return FileType.CSV
        elif filename_lower.endswith(('.xlsx', '.xls')):
            return FileType.EXCEL
        elif filename_lower.endswith('.pdf'):
            return FileType.PDF
        elif filename_lower.endswith(('.docx', '.doc')):
            return FileType.DOCX
        
        # Try to detect from content
        if content[:4] == b'%PDF':
            return FileType.PDF
        elif content[:4] == b'PK\x03\x04':  # ZIP signature (xlsx, docx)
            return FileType.EXCEL  # Could be either, default to Excel
        
        return FileType.UNKNOWN
    
    def detect_encoding(self, content: bytes) -> str:
        """Detect text encoding of file content"""
        result = chardet.detect(content)
        return result.get('encoding', 'utf-8') or 'utf-8'
    
    async def parse_file(self, filename: str, content: bytes) -> ParseResult:
        """
        Main entry point for parsing any supported file
        """
        file_type = self.detect_file_type(filename, content)
        
        try:
            if file_type == FileType.CSV:
                return await self.parse_csv(content)
            elif file_type == FileType.EXCEL:
                return await self.parse_excel(content)
            elif file_type == FileType.PDF:
                return await self.parse_pdf(content)
            elif file_type == FileType.DOCX:
                return await self.parse_docx(content)
            else:
                # Try CSV as fallback
                return await self.parse_csv(content)
        except Exception as e:
            traceback.print_exc()
            return ParseResult(
                success=False,
                students=[],
                file_type=file_type,
                columns_detected=[],
                mapping_confidence=0.0,
                warnings=[f"Error parsing file: {str(e)}"]
            )
    
    async def parse_csv(self, content: bytes) -> ParseResult:
        """Parse CSV file content"""
        warnings = []
        
        # Detect encoding
        encoding = self.detect_encoding(content)
        
        # Try different delimiters
        text_content = content.decode(encoding, errors='replace')
        
        # Detect delimiter
        delimiters = [',', ';', '\t', '|']
        best_delimiter = ','
        max_cols = 0
        
        for delim in delimiters:
            try:
                df_test = pd.read_csv(io.StringIO(text_content), sep=delim, nrows=5)
                if len(df_test.columns) > max_cols:
                    max_cols = len(df_test.columns)
                    best_delimiter = delim
            except:
                continue
        
        # Read full CSV
        try:
            df = pd.read_csv(io.StringIO(text_content), sep=best_delimiter)
        except Exception as e:
            warnings.append(f"CSV parsing with standard method failed: {str(e)}")
            # Try with more lenient settings
            df = pd.read_csv(
                io.StringIO(text_content), 
                sep=best_delimiter,
                on_bad_lines='skip',
                encoding_errors='replace'
            )
        
        return self._process_dataframe(df, FileType.CSV, warnings)
    
    async def parse_excel(self, content: bytes) -> ParseResult:
        """Parse Excel file content"""
        warnings = []
        
        try:
            # Read all sheets
            xlsx = pd.ExcelFile(io.BytesIO(content))
            sheets = xlsx.sheet_names
            
            # Find the best sheet (usually the one with student data)
            best_df = None
            best_score = 0
            
            for sheet in sheets:
                df = pd.read_excel(xlsx, sheet_name=sheet)
                if df.empty:
                    continue
                    
                # Score based on column matches
                mapping = ColumnMapper.map_columns(list(df.columns.astype(str)))
                score = ColumnMapper.calculate_confidence(mapping)
                
                if score > best_score:
                    best_score = score
                    best_df = df
            
            if best_df is None:
                best_df = pd.read_excel(io.BytesIO(content), sheet_name=0)
            
            if len(sheets) > 1:
                warnings.append(f"Multiple sheets detected: {sheets}. Using the one with best data match.")
            
            return self._process_dataframe(best_df, FileType.EXCEL, warnings)
            
        except Exception as e:
            return ParseResult(
                success=False,
                students=[],
                file_type=FileType.EXCEL,
                columns_detected=[],
                mapping_confidence=0.0,
                warnings=[f"Excel parsing error: {str(e)}"]
            )
    
    async def parse_pdf(self, content: bytes) -> ParseResult:
        """Parse PDF file content - extracts tables and text"""
        warnings = []
        extracted_data = []
        
        if not PDF_SUPPORT:
            return ParseResult(
                success=False,
                students=[],
                file_type=FileType.PDF,
                columns_detected=[],
                mapping_confidence=0.0,
                warnings=["PDF support not installed. Please install pdfplumber."]
            )
        
        try:
            import pdfplumber as pdfplumber_lib
            with pdfplumber_lib.open(io.BytesIO(content)) as pdf:
                all_tables = []
                all_text = []
                
                for page in pdf.pages:
                    # Extract tables
                    tables = page.extract_tables()
                    for table in tables:
                        if table and len(table) > 1:  # Has header + data
                            all_tables.append(table)
                    
                    # Extract text for context
                    text = page.extract_text()
                    if text:
                        all_text.append(text)
                
                # Try to convert tables to DataFrame
                if all_tables:
                    # Use the largest table
                    best_table = max(all_tables, key=lambda t: len(t) * len(t[0]) if t else 0)
                    
                    # First row as headers
                    headers = [str(h).strip() if h else f"Column_{i}" for i, h in enumerate(best_table[0])]
                    data = best_table[1:]
                    
                    df = pd.DataFrame(data, columns=headers)
                    return self._process_dataframe(df, FileType.PDF, warnings)
                
                # If no tables, try to extract from text using AI
                if all_text:
                    full_text = "\n".join(all_text)
                    return await self._parse_unstructured_text(full_text, FileType.PDF)
                
                return ParseResult(
                    success=False,
                    students=[],
                    file_type=FileType.PDF,
                    columns_detected=[],
                    mapping_confidence=0.0,
                    warnings=["No extractable tables or text found in PDF"]
                )
                
        except Exception as e:
            traceback.print_exc()
            return ParseResult(
                success=False,
                students=[],
                file_type=FileType.PDF,
                columns_detected=[],
                mapping_confidence=0.0,
                warnings=[f"PDF parsing error: {str(e)}"]
            )
    
    async def parse_docx(self, content: bytes) -> ParseResult:
        """Parse Word document content"""
        warnings = []
        
        if not DOCX_SUPPORT:
            return ParseResult(
                success=False,
                students=[],
                file_type=FileType.DOCX,
                columns_detected=[],
                mapping_confidence=0.0,
                warnings=["DOCX support not installed. Please install python-docx."]
            )
        
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(content))
            
            # Extract tables
            all_data = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                if table_data:
                    all_data.append(table_data)
            
            if all_data:
                # Use largest table
                best_table = max(all_data, key=lambda t: len(t) * len(t[0]) if t else 0)
                headers = best_table[0]
                data = best_table[1:]
                df = pd.DataFrame(data, columns=headers)
                return self._process_dataframe(df, FileType.DOCX, warnings)
            
            # Extract text if no tables
            full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            if full_text:
                return await self._parse_unstructured_text(full_text, FileType.DOCX)
            
            return ParseResult(
                success=False,
                students=[],
                file_type=FileType.DOCX,
                columns_detected=[],
                mapping_confidence=0.0,
                warnings=["No extractable tables or text found in document"]
            )
            
        except Exception as e:
            return ParseResult(
                success=False,
                students=[],
                file_type=FileType.DOCX,
                columns_detected=[],
                mapping_confidence=0.0,
                warnings=[f"DOCX parsing error: {str(e)}"]
            )
    
    def _process_dataframe(self, df: pd.DataFrame, file_type: FileType, warnings: List[str]) -> ParseResult:
        """Process a pandas DataFrame into ParseResult"""
        
        # Clean up column names
        df.columns = [str(c).strip() for c in df.columns]
        
        # Drop completely empty rows and columns
        df = df.dropna(how='all').dropna(axis=1, how='all')
        
        if df.empty:
            return ParseResult(
                success=False,
                students=[],
                file_type=file_type,
                columns_detected=[],
                mapping_confidence=0.0,
                warnings=warnings + ["No valid data found in file"]
            )
        
        columns = list(df.columns)
        mapping = ColumnMapper.map_columns(columns)
        confidence = ColumnMapper.calculate_confidence(mapping)
        
        # Extract students
        students = []
        
        for idx, row in df.iterrows():
            try:
                # Get name
                name = None
                if mapping['name']:
                    name = str(row[mapping['name']]).strip()
                
                if not name or name.lower() in ['nan', 'none', '']:
                    continue
                
                # Get quiz score (may need to average multiple columns)
                quiz_score = None
                if mapping['quiz_score']:
                    val = row[mapping['quiz_score']]
                    quiz_score = self._extract_numeric(val)
                else:
                    # Try to find any numeric columns that look like grades
                    numeric_cols = self._find_grade_columns(df, columns, row)
                    if numeric_cols:
                        quiz_score = np.mean(numeric_cols)
                
                # Get engagement
                engagement = None
                if mapping['engagement']:
                    engagement = self._extract_numeric(row[mapping['engagement']])
                
                # Get attendance
                attendance = None
                if mapping['attendance']:
                    attendance = self._extract_numeric(row[mapping['attendance']])
                
                # If no engagement but have attendance, use attendance as engagement proxy
                if engagement is None and attendance is not None:
                    # Check if there's a total days column for calculating percentage
                    total_days = None
                    for col in columns:
                        col_lower = col.lower()
                        if 'total' in col_lower and ('day' in col_lower or 'session' in col_lower or 'class' in col_lower):
                            total_days = self._extract_numeric(row[col])
                            break
                    
                    if total_days and total_days > 0:
                        # attendance is days present, calculate percentage
                        engagement = (attendance / total_days) * 100
                    elif attendance <= 100:
                        # attendance is already a percentage
                        engagement = attendance
                
                # Get weak topic
                weak_topic = None
                if mapping['weak_topic']:
                    weak_topic = str(row[mapping['weak_topic']]).strip()
                
                # Extract all grades
                grades = {}
                for col in columns:
                    val = row[col]
                    if self._looks_like_grade_column(col) and pd.notna(val):
                        numeric_val = self._extract_numeric(val)
                        if numeric_val is not None:
                            grades[col] = numeric_val
                
                # Create raw data dict
                raw_data = {col: str(row[col]) if pd.notna(row[col]) else None for col in columns}
                
                student = ParsedStudent(
                    name=name,
                    engagement_score=float(engagement) if engagement is not None else None,
                    avg_quiz_score=float(quiz_score) if quiz_score is not None else None,
                    attendance=float(attendance) if attendance is not None else None,
                    weakest_topic=weak_topic,
                    grades=grades if grades else None,
                    raw_data=raw_data
                )
                students.append(student)
                
            except Exception as e:
                warnings.append(f"Error processing row {idx}: {str(e)}")
                continue
        
        # Generate preview
        preview = df.head(5).to_string() if len(df) > 0 else None
        
        return ParseResult(
            success=len(students) > 0,
            students=students,
            file_type=file_type,
            columns_detected=columns,
            mapping_confidence=confidence,
            warnings=warnings,
            raw_preview=preview
        )
    
    def _extract_numeric(self, value: Any) -> Optional[float]:
        """Extract numeric value from various formats"""
        if pd.isna(value):
            return None
        
        if isinstance(value, (int, float)):
            return float(value)
        
        str_val = str(value).strip()
        
        # Handle percentage format
        if '%' in str_val:
            str_val = str_val.replace('%', '').strip()
        
        # Handle fraction format (e.g., "85/100")
        if '/' in str_val:
            parts = str_val.split('/')
            if len(parts) == 2:
                try:
                    return (float(parts[0]) / float(parts[1])) * 100
                except:
                    pass
        
        # Handle letter grades
        grade_map = {
            'A+': 97, 'A': 93, 'A-': 90,
            'B+': 87, 'B': 83, 'B-': 80,
            'C+': 77, 'C': 73, 'C-': 70,
            'D+': 67, 'D': 63, 'D-': 60,
            'F': 50
        }
        if str_val.upper() in grade_map:
            return float(grade_map[str_val.upper()])
        
        # Try direct conversion
        try:
            return float(str_val)
        except:
            return None
    
    def _looks_like_grade_column(self, col_name: str) -> bool:
        """Check if column name looks like it contains grades"""
        patterns = [
            r'quiz', r'test', r'exam', r'grade', r'score', r'mark',
            r'hw', r'homework', r'assignment', r'midterm', r'final',
            r'q\d+', r'quiz\s*\d+', r'test\s*\d+', r'unit\s*\d+'
        ]
        col_lower = col_name.lower()
        return any(re.search(p, col_lower) for p in patterns)
    
    def _find_grade_columns(self, df: pd.DataFrame, columns: List[str], row: pd.Series) -> List[float]:
        """Find and extract values from columns that look like grades"""
        grades = []
        for col in columns:
            if self._looks_like_grade_column(col):
                val = self._extract_numeric(row[col])
                if val is not None and 0 <= val <= 100:
                    grades.append(val)
        return grades
    
    async def _parse_unstructured_text(self, text: str, file_type: FileType) -> ParseResult:
        """
        Parse unstructured text using LLM to extract student data
        Fallback for documents without clear table structure
        """
        warnings = ["Document parsed as unstructured text. Results may need verification."]
        
        # If we have HuggingFace service, use it for intelligent extraction
        if self.hf_service:
            try:
                return await self._llm_extract_students(text, file_type, warnings)
            except Exception as e:
                warnings.append(f"AI extraction failed: {str(e)}")
        
        # Fallback: try to extract using patterns
        students = self._pattern_extract_students(text)
        
        return ParseResult(
            success=len(students) > 0,
            students=students,
            file_type=file_type,
            columns_detected=[],
            mapping_confidence=0.3 if students else 0.0,
            warnings=warnings,
            raw_preview=text[:500] if text else None
        )
    
    async def _llm_extract_students(self, text: str, file_type: FileType, warnings: List[str]) -> ParseResult:
        """Use LLM to extract student data from unstructured text"""
        
        # Truncate text to fit in context window
        max_text = text[:4000]
        
        prompt = f"""Analyze this educational document and extract student information.
Return ONLY a valid JSON array with student objects.

Each student object should have these fields (use null if not found):
- name: student's full name
- quiz_score: numeric grade/score (0-100)
- engagement: participation/engagement score (0-100)
- attendance: attendance percentage (0-100)
- weak_topic: area needing improvement

Document content:
{max_text}

Respond with ONLY the JSON array, no other text:"""

        messages = [
            {"role": "system", "content": "You are a data extraction assistant. Extract student data from documents and return only valid JSON."},
            {"role": "user", "content": prompt}
        ]
        
        if self.hf_service is None:
            warnings.append("AI service not available for text extraction")
            return ParseResult(
                success=False,
                students=[],
                file_type=file_type,
                columns_detected=[],
                mapping_confidence=0.0,
                warnings=warnings,
                raw_preview=text[:500] if text else None
            )
        
        try:
            model_name = "Qwen/Qwen2.5-3B-Instruct"
            if hasattr(self.hf_service, 'client') and self.hf_service.client is not None:
                model_name = getattr(self.hf_service.client, 'model', model_name)
            result = await self.hf_service.query_chat(
                model_name,
                messages,
                max_tokens=2000
            )
            
            response_text = result.get("generated_text", "")
            
            # Try to extract JSON from response
            json_match = re.search(r'\[[\s\S]*?\]', response_text)
            if json_match:
                students_data = json.loads(json_match.group())
                students = []
                
                for s in students_data:
                    if isinstance(s, dict) and s.get('name'):
                        quiz_score = s.get('quiz_score')
                        engagement = s.get('engagement')
                        attendance = s.get('attendance')
                        students.append(ParsedStudent(
                            name=str(s.get('name')),
                            avg_quiz_score=float(quiz_score) if quiz_score is not None else None,
                            engagement_score=float(engagement) if engagement is not None else None,
                            attendance=float(attendance) if attendance is not None else None,
                            weakest_topic=s.get('weak_topic')
                        ))
                
                return ParseResult(
                    success=len(students) > 0,
                    students=students,
                    file_type=file_type,
                    columns_detected=['name', 'quiz_score', 'engagement', 'attendance', 'weak_topic'],
                    mapping_confidence=0.7,
                    warnings=warnings + ["Data extracted using AI. Please verify accuracy."],
                    raw_preview=text[:500]
                )
        except json.JSONDecodeError:
            warnings.append("AI response was not valid JSON")
        except Exception as e:
            warnings.append(f"AI extraction error: {str(e)}")
        
        # Fallback to pattern extraction
        students = self._pattern_extract_students(text)
        return ParseResult(
            success=len(students) > 0,
            students=students,
            file_type=file_type,
            columns_detected=[],
            mapping_confidence=0.3,
            warnings=warnings,
            raw_preview=text[:500]
        )
    
    def _pattern_extract_students(self, text: str) -> List[ParsedStudent]:
        """Extract students using regex patterns - fallback method"""
        students = []
        
        # Pattern for "Name: X" or "Student: X" followed by grades
        name_patterns = [
            r'(?:student|name|learner)[\s:]+([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+)*)',
            r'^([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+)+)[\s:,]+\d',
        ]
        
        lines = text.split('\n')
        for line in lines:
            for pattern in name_patterns:
                match = re.search(pattern, line, re.IGNORECASE | re.MULTILINE)
                if match:
                    name = match.group(1).strip()
                    
                    # Try to extract grades from same line
                    numbers = re.findall(r'\b(\d{1,3}(?:\.\d+)?)\b', line)
                    grades = [float(n) for n in numbers if 0 <= float(n) <= 100]
                    
                    avg_score = float(np.mean(grades)) if grades else None
                    
                    students.append(ParsedStudent(
                        name=name,
                        avg_quiz_score=avg_score,
                        raw_data={"original_line": line}
                    ))
                    break
        
        return students


class CourseMaterialParser:
    """Parser for course materials (syllabus, lesson plans, curriculum)"""
    
    def __init__(self, hf_service=None):
        self.hf_service = hf_service
    
    async def parse_course_material(self, filename: str, content: bytes) -> Dict[str, Any]:
        """Parse course material and extract relevant information"""
        
        # Detect file type and extract text
        text = await self._extract_text(filename, content)
        
        if not text:
            return {
                "success": False,
                "error": "Could not extract text from file"
            }
        
        # Extract course information
        course_info = await self._analyze_course_content(text)
        
        return {
            "success": True,
            "filename": filename,
            "courseInfo": course_info,
            "textPreview": text[:1000]
        }
    
    async def _extract_text(self, filename: str, content: bytes) -> Optional[str]:
        """Extract text from various file formats"""
        filename_lower = filename.lower()
        
        if filename_lower.endswith('.pdf'):
            if PDF_SUPPORT:
                try:
                    import pdfplumber as pdfplumber_lib
                    with pdfplumber_lib.open(io.BytesIO(content)) as pdf:
                        return "\n".join(page.extract_text() or "" for page in pdf.pages)
                except:
                    pass
        
        elif filename_lower.endswith(('.docx', '.doc')):
            if DOCX_SUPPORT:
                try:
                    from docx import Document as DocxDocument
                    doc = DocxDocument(io.BytesIO(content))
                    return "\n".join(p.text for p in doc.paragraphs)
                except:
                    pass
        
        elif filename_lower.endswith('.txt'):
            try:
                encoding = chardet.detect(content).get('encoding', 'utf-8')
                return content.decode(encoding or 'utf-8')
            except:
                pass
        
        return None
    
    async def _analyze_course_content(self, text: str) -> Dict[str, Any]:
        """Analyze course content and extract topics, structure, etc."""
        
        # Extract topics using patterns
        topics = []
        topic_patterns = [
            r'(?:topic|unit|chapter|module|lesson)[\s:#]+(.+?)(?:\n|$)',
            r'(?:\d+\.|\-|\•)\s*(.+?)(?:\n|$)'
        ]
        
        for pattern in topic_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            topics.extend([m.strip() for m in matches if len(m.strip()) > 3])
        
        # Extract dates
        dates = re.findall(
            r'\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s+\d{1,2}(?:,?\s+\d{4})?\b',
            text, re.IGNORECASE
        )
        
        # Detect math topics
        math_keywords = [
            'algebra', 'calculus', 'geometry', 'trigonometry', 'statistics',
            'derivative', 'integral', 'equation', 'function', 'graph',
            'probability', 'linear', 'quadratic', 'polynomial'
        ]
        detected_topics = [kw for kw in math_keywords if kw in text.lower()]
        
        return {
            "detectedTopics": list(set(detected_topics)),
            "outlineTopics": list(set(topics[:20])),  # Limit to 20
            "dates": list(set(dates[:10])),  # Limit to 10
            "wordCount": len(text.split()),
            "hasAssessments": any(w in text.lower() for w in ['quiz', 'test', 'exam', 'assessment']),
            "hasObjectives": any(w in text.lower() for w in ['objective', 'goal', 'outcome', 'learning'])
        }


# Create singleton instance
document_parser = DocumentParser()
course_parser = CourseMaterialParser()
